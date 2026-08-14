from pathlib import Path

from docx import Document

from processing.cleaner import TextCleaner

from loaders.base_loader import BaseLoader


class DOCXLoader(BaseLoader):

    def __init__(self, file_path):

        self.file_path = Path(file_path)

    def load(self):

        if not self.file_path.exists():

            raise FileNotFoundError(
                f"{self.file_path} not found."
            )

        doc = Document(
            self.file_path
        )

        parts = [
            paragraph.text
            for paragraph in doc.paragraphs
        ]

        # Tables are common in specs/pricing/comparison docs and were
        # previously dropped entirely.
        for table in doc.tables:

            for row in table.rows:

                row_text = " | ".join(
                    cell.text for cell in row.cells
                )

                if row_text.strip():
                    parts.append(row_text)

        text = "\n".join(parts)

        return [

            {

                "page_number": 1,

                "text": TextCleaner.clean(text),

                "metadata": {

                    "source": self.file_path.name,

                    "page": 1,

                    "file_type": "docx"

                }

            }

        ]
